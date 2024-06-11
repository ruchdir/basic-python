import os
import barcode
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Inches

# List of strings to convert to barcodes
data_list = ["123456", "654321", "112233", "445566"]

# Function to generate barcode images
def generate_barcodes(data_list, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    filenames = []
    for data in data_list:
        try:
            # Specify Code-128 format and use ImageWriter to generate an image file
            code128 = barcode.get_barcode_class('code128')
            barcode_obj = code128(data, writer=ImageWriter())

            # Save the barcode as an image file in the output directory
            filename = os.path.join(output_dir, f'barcode_{data}')
            barcode_obj.save(filename)
            filename_with_extension = f'{filename}.png'
            
            if os.path.exists(filename_with_extension):
                filenames.append(filename_with_extension)
                print(f"Barcode for {data} saved as {filename_with_extension}")
            else:
                print(f"Error: Barcode for {data} was not saved.")
        
        except Exception as e:
            print(f"An error occurred while generating barcode for {data}: {e}")
    
    return filenames

# Generate the barcodes
output_dir = os.path.abspath('barcodes')  # Use absolute path for the output directory
barcode_files = generate_barcodes(data_list, output_dir)

# Check if any barcodes were generated
if not barcode_files:
    print("No barcodes were generated.")
else:
    # Create a new Word document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Barcode List', level=1)

    # Insert each barcode image into the document
    for barcode_file in barcode_files:
        doc.add_paragraph(os.path.basename(barcode_file))  # Add the filename as a paragraph
        doc.add_picture(barcode_file, width=Inches(2))  # Add the barcode image

    # Save the document
    doc_filename = 'barcodes.docx'
    doc.save(doc_filename)
    print(f"Word document saved as {doc_filename}")
